from __future__ import annotations

import json
import sys
from pathlib import Path


def extract_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    chunks: list[str] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        text = text.strip()
        if text:
          chunks.append(text)

    return "\n\n".join(chunks).strip()


def extract_docx(path: Path) -> str:
    from docx import Document

    document = Document(str(path))
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]

    table_chunks: list[str] = []
    for table in document.tables:
        rows: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                rows.append(" | ".join(cells))
        if rows:
            table_chunks.append("\n".join(rows))

    chunks = paragraphs + table_chunks
    return "\n\n".join(chunks).strip()


def main() -> int:
    if len(sys.argv) < 2:
        print(json.dumps({"ok": False, "error": "Missing document path"}))
        return 1

    path = Path(sys.argv[1]).expanduser().resolve()
    if not path.exists():
        print(json.dumps({"ok": False, "error": f"File not found: {path}"}))
        return 1

    suffix = path.suffix.lower()

    try:
        if suffix == ".pdf":
            content = extract_pdf(path)
        elif suffix == ".docx":
            content = extract_docx(path)
        else:
            print(json.dumps({"ok": False, "error": f"Unsupported extension: {suffix}"}))
            return 1
    except Exception as exc:  # noqa: BLE001
        print(json.dumps({"ok": False, "error": str(exc)}))
        return 1

    print(
        json.dumps(
            {
                "ok": True,
                "content": content,
                "characterCount": len(content),
            },
            ensure_ascii=False,
        )
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
